import os
import re
from typing import List, Dict, Any, Optional
from pathlib import Path
import hashlib
from datetime import datetime
import json

# Importaciones para diferentes formatos
try:
    import pypdf
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    
from docx import Document as DocxDocument
import pandas as pd
from bs4 import BeautifulSoup
import chardet
from unidecode import unidecode

from ..utils.logger import get_logger
from ..config.settings import settings

logger = get_logger(__name__)

class Document:
    """Clase para representar un documento"""
    def __init__(
        self,
        content: str,
        metadata: Dict[str, Any],
        doc_id: Optional[str] = None
    ):
        self.content = content
        self.metadata = metadata
        self.doc_id = doc_id or self._generate_id()
        
    def _generate_id(self) -> str:
        """Genera un ID único para el documento"""
        content_hash = hashlib.md5(self.content.encode()).hexdigest()[:8]
        return f"{self.metadata.get('filename', 'unknown')}_{content_hash}"
        
    def to_dict(self) -> Dict[str, Any]:
        """Convierte el documento a diccionario"""
        return {
            "doc_id": self.doc_id,
            "content": self.content,
            "metadata": self.metadata
        }

class DocumentLoader:
    """Cargador avanzado de documentos con soporte para múltiples formatos"""
    
    SUPPORTED_EXTENSIONS = {
        '.pdf': 'load_pdf',
        '.docx': 'load_docx',
        '.doc': 'load_doc',
        '.txt': 'load_text',
        '.html': 'load_html',
        '.csv': 'load_csv',
        '.xlsx': 'load_excel',
        '.xls': 'load_excel'
    }
    
    def __init__(self):
        self.loaded_documents = []
        self.extraction_stats = {
            "total_files": 0,
            "successful": 0,
            "failed": 0,
            "total_chars": 0,
            "extraction_time": 0
        }
        
    def load_directory(self, directory_path: str) -> List[Document]:
        """Carga todos los documentos de un directorio"""
        logger.info(f"Cargando documentos desde: {directory_path}")
        
        documents = []
        start_time = datetime.now()
        
        for file_path in Path(directory_path).rglob("*"):
            if file_path.is_file():
                try:
                    doc = self.load_document(str(file_path))
                    if doc:
                        documents.append(doc)
                        self.extraction_stats["successful"] += 1
                except Exception as e:
                    logger.error(f"Error cargando {file_path}: {str(e)}")
                    self.extraction_stats["failed"] += 1
                    
                self.extraction_stats["total_files"] += 1
                
        self.extraction_stats["extraction_time"] = (datetime.now() - start_time).total_seconds()
        self.loaded_documents.extend(documents)
        
        logger.info(f"Cargados {len(documents)} documentos correctamente")
        logger.info(f"Estadísticas: {self.extraction_stats}")
        
        return documents
        
    def load_document(self, file_path: str) -> Optional[Document]:
        """Carga un documento individual"""
        file_path = Path(file_path)
        
        if not file_path.exists():
            logger.error(f"El archivo no existe: {file_path}")
            return None
            
        extension = file_path.suffix.lower()
        
        if extension not in self.SUPPORTED_EXTENSIONS:
            logger.warning(f"Formato no soportado: {extension}")
            return None
            
        # Obtener metadata básica
        metadata = self._extract_metadata(file_path)
        
        # Cargar contenido según el tipo
        loader_method = getattr(self, self.SUPPORTED_EXTENSIONS[extension])
        content = loader_method(file_path)
        
        if content:
            # Limpiar y normalizar el contenido
            content = self._clean_content(content)
            self.extraction_stats["total_chars"] += len(content)
            
            # Extraer metadata adicional del contenido
            metadata.update(self._extract_content_metadata(content))
            
            return Document(content=content, metadata=metadata)
            
        return None
        
    def load_pdf(self, file_path: Path) -> str:
        """Carga un archivo PDF"""
        logger.debug(f"Cargando PDF: {file_path}")
        
        if not PDF_AVAILABLE:
            logger.error("pypdf no está instalado. Instálalo con: pip install pypdf")
            return ""
        
        text_parts = []
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = pypdf.PdfReader(file)  # Cambiado de PyPDF2.PdfReader a pypdf.PdfReader
                
                for page_num, page in enumerate(pdf_reader.pages):
                    text = page.extract_text()
                    if text.strip():
                        # Agregar marcador de página para referencia
                        text_parts.append(f"[Página {page_num + 1}]\n{text}")
                        
            return "\n\n".join(text_parts)
            
        except Exception as e:
            logger.error(f"Error leyendo PDF {file_path}: {str(e)}")
            return ""
            
    def load_docx(self, file_path: Path) -> str:
        """Carga un archivo DOCX"""
        logger.debug(f"Cargando DOCX: {file_path}")
        
        try:
            doc = DocxDocument(file_path)
            
            # Extraer texto de párrafos
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
                    
            # Extraer texto de tablas
            tables_text = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    if any(row_data):
                        table_data.append(" | ".join(row_data))
                        
                if table_data:
                    tables_text.append("\n".join(table_data))
                    
            # Combinar todo
            all_text = "\n\n".join(paragraphs)
            if tables_text:
                all_text += "\n\n[TABLAS]\n" + "\n\n".join(tables_text)
                
            return all_text
            
        except Exception as e:
            logger.error(f"Error leyendo DOCX {file_path}: {str(e)}")
            return ""
            
    def load_text(self, file_path: Path) -> str:
        """Carga un archivo de texto plano"""
        logger.debug(f"Cargando TXT: {file_path}")
        
        try:
            # Detectar encoding
            with open(file_path, 'rb') as file:
                raw_data = file.read()
                result = chardet.detect(raw_data)
                encoding = result['encoding'] or 'utf-8'
                
            # Leer con el encoding detectado
            with open(file_path, 'r', encoding=encoding) as file:
                return file.read()
                
        except Exception as e:
            logger.error(f"Error leyendo texto {file_path}: {str(e)}")
            return ""
            
    def load_html(self, file_path: Path) -> str:
        """Carga un archivo HTML y extrae el texto"""
        logger.debug(f"Cargando HTML: {file_path}")
        
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                soup = BeautifulSoup(file.read(), 'html.parser')
                
                # Eliminar scripts y estilos
                for script in soup(['script', 'style']):
                    script.decompose()
                    
                # Obtener texto
                text = soup.get_text()
                
                # Limpiar líneas vacías múltiples
                lines = (line.strip() for line in text.splitlines())
                chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
                text = '\n'.join(chunk for chunk in chunks if chunk)
                
                return text
                
        except Exception as e:
            logger.error(f"Error leyendo HTML {file_path}: {str(e)}")
            return ""
            
    def load_excel(self, file_path: Path) -> str:
        """Carga un archivo Excel"""
        logger.debug(f"Cargando Excel: {file_path}")
        
        try:
            # Leer todas las hojas
            excel_file = pd.ExcelFile(file_path)
            
            all_text = []
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                
                # Convertir a texto
                sheet_text = f"[Hoja: {sheet_name}]\n"
                sheet_text += df.to_string(index=False)
                all_text.append(sheet_text)
                
            return "\n\n".join(all_text)
            
        except Exception as e:
            logger.error(f"Error leyendo Excel {file_path}: {str(e)}")
            return ""
            
    def load_csv(self, file_path: Path) -> str:
        """Carga un archivo CSV"""
        logger.debug(f"Cargando CSV: {file_path}")
        
        try:
            df = pd.read_csv(file_path)
            return df.to_string(index=False)
            
        except Exception as e:
            logger.error(f"Error leyendo CSV {file_path}: {str(e)}")
            return ""
            
    def load_doc(self, file_path: Path) -> str:
        """Carga un archivo DOC (formato antiguo de Word)"""
        # Por ahora, intentamos como texto plano
        logger.warning(f"Formato .doc no totalmente soportado, intentando como texto: {file_path}")
        return self.load_text(file_path)
        
    def _extract_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extrae metadata básica del archivo"""
        stat = file_path.stat()
        
        return {
            "filename": file_path.name,
            "filepath": str(file_path),
            "extension": file_path.suffix,
            "size_bytes": stat.st_size,
            "created_at": datetime.fromtimestamp(stat.st_ctime).isoformat(),
            "modified_at": datetime.fromtimestamp(stat.st_mtime).isoformat(),
            "loaded_at": datetime.now().isoformat()
        }
        
    def _extract_content_metadata(self, content: str) -> Dict[str, Any]:
        """Extrae metadata adicional del contenido"""
        metadata = {
            "char_count": len(content),
            "word_count": len(content.split()),
            "line_count": len(content.splitlines())
        }
        
        # Detectar tipo de contrato
        contract_type = self._detect_contract_type(content)
        if contract_type:
            metadata["contract_type"] = contract_type
            
        # Extraer fechas
        dates = self._extract_dates(content)
        if dates:
            metadata["dates_found"] = dates
            
        # Extraer partes del contrato
        parties = self._extract_parties(content)
        if parties:
            metadata["parties"] = parties
            
        return metadata
        
    # def _detect_contract_type(self, content: str) -> Optional[str]:
    #     """Detecta el tipo de contrato basándose en palabras clave"""
    #     content_lower = content.lower()
        
    #     contract_types = {
    #         "arrendamiento": ["arrendamiento", "alquiler", "renta", "inquilino", "arrendador"],
    #         "gestión": ["gestión hotelera", "management", "operación hotelera", "gestor"],
    #         "préstamo": ["préstamo", "crédito", "financiación", "prestamista", "deudor"],
    #         "compraventa": ["compraventa", "compra", "venta", "vendedor", "comprador"],
    #         "servicios": ["servicios", "prestación", "proveedor", "contratista"]
    #     }
        
    #     for contract_type, keywords in contract_types.items():
    #         if any(keyword in content_lower for keyword in keywords):
    #             return contract_type
                
    #     return None

    def _detect_contract_type(self, content: str) -> Optional[str]:
        """Detecta el tipo de contrato basándose en palabras clave"""
        content_lower = content.lower()
        
        # IMPORTANTE: El orden importa - buscar primero los más específicos
        contract_types = {
            "franquicia": ["franquicia", "franquiciador", "franquiciado", "royalty", "fee de marketing", "occidental hotels & resorts"],
            "mantenimiento": ["mantenimiento", "servicios de mantenimiento", "preventivo", "correctivo", "technical maintenance"],
            "gestión": ["gestión hotelera", "management", "operación hotelera", "gestor", "barceló hotel group"],
            "arrendamiento": ["arrendamiento", "alquiler", "renta", "inquilino", "arrendador", "local de negocio"],
            "préstamo": ["préstamo", "crédito", "financiación", "prestamista", "deudor"],
            "compraventa": ["compraventa", "compra", "venta", "vendedor", "comprador"],
            "servicios": ["servicios", "prestación", "proveedor", "contratista"]
        }
        
        # Buscar en el título o primeras líneas (más confiable)
        first_lines = content[:1000].lower()
        
        # Búsqueda estricta en el título/inicio
        for contract_type, keywords in contract_types.items():
            for keyword in keywords:
                # Buscar coincidencias más específicas primero
                if contract_type == "franquicia" and "contrato de franquicia" in first_lines:
                    return "franquicia"
                elif contract_type == "mantenimiento" and "servicios de mantenimiento" in first_lines:
                    return "mantenimiento"
                elif keyword in first_lines:
                    return contract_type
        
        # Si no encontramos en el título, buscar en todo el documento
        # pero con un scoring para evitar falsos positivos
        scores = {}
        for contract_type, keywords in contract_types.items():
            score = 0
            for keyword in keywords:
                if keyword in content_lower:
                    # Dar más peso a palabras más específicas
                    if len(keyword) > 10:  # Palabras largas son más específicas
                        score += 2
                    else:
                        score += 1
            if score > 0:
                scores[contract_type] = score
        
        # Retornar el tipo con mayor score
        if scores:
            return max(scores.items(), key=lambda x: x[1])[0]
            
        return "otros" 
        
    def _extract_dates(self, content: str) -> List[str]:
        """Extrae fechas del contenido"""
        # Patrones comunes de fecha en español
        date_patterns = [
            r'\d{1,2}\s+de\s+\w+\s+de\s+\d{4}',
            r'\d{1,2}/\d{1,2}/\d{4}',
            r'\d{1,2}-\d{1,2}-\d{4}',
            r'\d{4}-\d{1,2}-\d{1,2}'
        ]
        
        dates = []
        for pattern in date_patterns:
            matches = re.findall(pattern, content, re.IGNORECASE)
            dates.extend(matches)
            
        return list(set(dates))[:10]  # Limitar a 10 fechas únicas
        
    def _extract_parties(self, content: str) -> List[str]:
        """Extrae las partes mencionadas en el contrato"""
        # Buscar después de "PARTES:" o similar
        patterns = [
            r'(?:PARTES:|entre|ENTRE)\s*([^,\n]+)\s*(?:y|Y)\s*([^,\n]+)',
            r'(?:El|La)\s+(?:ARRENDADOR|ARRENDATARIO|PRESTAMISTA|PRESTATARIO):\s*([^\n]+)',
            r'(?:PRIMERA|SEGUNDA)\.?-?\s*([^\n]+)'
        ]
        
        parties = []
        for pattern in patterns:
            matches = re.findall(pattern, content, re.IGNORECASE)
            for match in matches:
                if isinstance(match, tuple):
                    parties.extend(match)
                else:
                    parties.append(match)
                    
        # Limpiar y filtrar
        parties = [p.strip() for p in parties if len(p.strip()) > 3]
        return list(set(parties))[:5]  # Limitar a 5 partes únicas
        
    def _clean_content(self, content: str) -> str:
        """Limpia y normaliza el contenido"""
        # Eliminar caracteres de control
        content = ''.join(char for char in content if ord(char) >= 32 or char in '\n\t')
        
        # Normalizar espacios en blanco
        content = re.sub(r'\s+', ' ', content)
        content = re.sub(r'\n\s*\n', '\n\n', content)
        
        # Eliminar espacios al inicio y final
        content = content.strip()
        
        return content
        
    def save_extraction_report(self, output_path: str):
        """Guarda un reporte de la extracción"""
        report = {
            "extraction_stats": self.extraction_stats,
            "documents": [
                {
                    "doc_id": doc.doc_id,
                    "filename": doc.metadata.get("filename"),
                    "char_count": doc.metadata.get("char_count", 0),
                    "contract_type": doc.metadata.get("contract_type", "unknown")
                }
                for doc in self.loaded_documents
            ]
        }
        
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(report, f, indent=2, ensure_ascii=False)
            
        logger.info(f"Informe de extracción guardado en: {output_path}")